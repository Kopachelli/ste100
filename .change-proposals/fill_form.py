"""Fill the ASD-STE100 change form.

Writes the proposal text into the empty paragraphs under each heading of the
form's single table. Row 7 (the STEMG assessment) is left untouched.
"""
import copy
import re
import sys

import docx

DOCX = sys.argv[1]
FIELDS = sys.argv[2]

# ---- read the field text -------------------------------------------------
raw = open(FIELDS, encoding="utf-8").read()
parts = re.split(r"^FIELD (\d)\n", raw, flags=re.M)
fields = {}
for i in range(1, len(parts), 2):
    fields[int(parts[i])] = parts[i + 1].strip("\n")

ATTRIBUTION = {
    "Sent by:": "Khristian Kopachelli",
    "Company or organization:": "Independent",
    "Email:": "kopachelli@gmail.com",
}
DATE = "2026-08-31"

doc = docx.Document(DOCX)
table = doc.tables[0]


def style_from(cell):
    """Return a run to clone formatting from: the heading run of this cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            if run.text.strip():
                return run
    return None


def write_into(cell, text, bold=False):
    """Put text into the empty paragraphs under the heading of this cell."""
    template = style_from(cell)
    empties = [p for p in cell.paragraphs if not p.text.strip()]
    lines = text.split("\n")
    # Reuse the empty paragraphs, then add more as needed.
    for index, line in enumerate(lines):
        if index < len(empties):
            para = empties[index]
        else:
            para = cell.add_paragraph()
            if empties:
                para.paragraph_format.left_indent = empties[0].paragraph_format.left_indent
                para.style = empties[0].style
        run = para.add_run(line)
        if template is not None:
            run.font.name = template.font.name
            run.font.size = template.font.size
        run.bold = bold


# Rows 1 thru 5 hold the five content fields, in order.
for row_index, field_number in ((1, 1), (2, 2), (3, 3), (4, 4), (5, 5)):
    cell = table.rows[row_index].cells[0]
    heading = cell.paragraphs[0].text.strip()
    print("row %d -> field %d  [%s...]" % (row_index, field_number, heading[:44]))
    write_into(cell, fields[field_number])

# Row 6: attribution on the left, date on the right.
left = table.rows[6].cells[0]
for para in left.paragraphs:
    label = para.text.strip()
    if label in ATTRIBUTION:
        template = para.runs[-1] if para.runs else None
        run = para.add_run(" " + ATTRIBUTION[label])
        if template is not None:
            run.font.name = template.font.name
            run.font.size = template.font.size
        print("row 6 -> %s %s" % (label, ATTRIBUTION[label]))

right = table.rows[6].cells[1]
for para in right.paragraphs:
    if para.text.strip() == "Date:":
        template = para.runs[-1] if para.runs else None
        run = para.add_run(" " + DATE)
        if template is not None:
            run.font.name = template.font.name
            run.font.size = template.font.size
        print("row 6 -> Date: %s" % DATE)

# Row 7 is the STEMG assessment. Do not write anything into it.
print("row 7 -> left blank for the STEMG")

doc.save(DOCX)
print("saved %s" % DOCX)
